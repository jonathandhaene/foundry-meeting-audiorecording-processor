"""
FastAPI application for the Meeting Audio Transcription service.

Provides REST API endpoints for uploading audio files and running
transcription with different methods (Azure Speech, Whisper).
"""

import os
import json
import logging
import uuid
import threading
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path
from typing import Dict, Any, Optional, List
from datetime import datetime, timezone
from enum import Enum

from fastapi import FastAPI, File, UploadFile, HTTPException, BackgroundTasks, Form, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field

from ..transcription.transcriber import AzureSpeechTranscriber
from ..transcription.whisper_transcriber import WhisperTranscriber
from ..transcription.hf_transcriber import HuggingFaceTranscriber
from ..audio.preprocessor import AudioPreprocessor
from ..nlp.analyzer import ContentAnalyzer
from ..utils.config import ConfigManager
from ..utils.logging import setup_logging

# Configure logging so ALL module loggers produce visible output
setup_logging(level="INFO")

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Persistent job storage (file-backed, survives container restarts)
# ---------------------------------------------------------------------------


class PersistentJobStore:
    """Thread-safe, file-backed job storage."""

    def __init__(self, path: str = "/home/meeting_transcription/jobs.json"):
        self._path = Path(path)
        self._path.parent.mkdir(parents=True, exist_ok=True)
        self._lock = threading.Lock()
        self._data: Dict[str, Dict[str, Any]] = self._load()

    def _load(self) -> Dict[str, Dict[str, Any]]:
        if self._path.exists():
            try:
                return json.loads(self._path.read_text())
            except Exception as e:
                logger.warning(f"Failed to load jobs from {self._path}: {e}")
        return {}

    def _save(self) -> None:
        try:
            self._path.write_text(json.dumps(self._data, default=str))
        except Exception as e:
            logger.warning(f"Failed to persist jobs to {self._path}: {e}")

    def __contains__(self, key: str) -> bool:
        with self._lock:
            return key in self._data

    def __getitem__(self, key: str) -> Dict[str, Any]:
        with self._lock:
            return self._data[key]

    def __setitem__(self, key: str, value: Dict[str, Any]) -> None:
        with self._lock:
            self._data[key] = value
            self._save()

    def __delitem__(self, key: str) -> None:
        with self._lock:
            del self._data[key]
            self._save()

    def values(self):
        with self._lock:
            return list(self._data.values())

    def clear(self) -> None:
        """Remove all jobs (used mainly in tests)."""
        with self._lock:
            self._data.clear()
            self._save()

    def update_field(self, key: str, field: str, value: Any) -> None:
        """Update a single field in a job and persist."""
        with self._lock:
            if key in self._data:
                self._data[key][field] = value
                self._save()

    def update_fields(self, key: str, fields: Dict[str, Any]) -> None:
        """Update multiple fields in a job and persist (single write)."""
        with self._lock:
            if key in self._data:
                self._data[key].update(fields)
                self._save()


# Initialize FastAPI app
app = FastAPI(
    title="Meeting Audio Transcription API",
    description="Upload audio files and transcribe using Azure Speech Services or Whisper",
    version="1.0.0",
)

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Configure appropriately for production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Job storage (file-backed, survives container restarts)
jobs_db = PersistentJobStore()

# Temporary file storage directory (use persistent /home/ mount on Azure App Service)
AUDIO_DIR = Path("/home/meeting_transcription/audio")
AUDIO_DIR.mkdir(parents=True, exist_ok=True)

# Export constants
MAX_KEY_PHRASES_EXPORT = 20  # Maximum number of key phrases to include in exports
MAX_SEGMENTS_TIMELINE = 20  # Maximum number of segments to show in audio timeline


class TranscriptionMethod(str, Enum):
    """Available transcription methods."""

    AZURE = "azure"
    WHISPER_LOCAL = "whisper_local"
    WHISPER_API = "whisper_api"
    HUGGINGFACE = "huggingface"


class JobStatus(str, Enum):
    """Job status enumeration."""

    PENDING = "pending"
    PROCESSING = "processing"
    COMPLETED = "completed"
    FAILED = "failed"


class TranscriptionRequest(BaseModel):
    """Request model for transcription."""

    method: TranscriptionMethod = Field(default=TranscriptionMethod.AZURE, description="Transcription method to use")
    language: Optional[str] = Field(default=None, description="Language code (e.g., 'en-US', 'es-ES') or None for auto-detect")
    enable_diarization: bool = Field(default=True, description="Enable speaker diarization (Azure only)")
    chunk_size: Optional[int] = Field(default=None, description="Audio chunk size in seconds for large files")
    whisper_model: str = Field(default="base", description="Whisper model size (tiny, base, small, medium, large)")
    enable_nlp: bool = Field(default=True, description="Enable NLP analysis (key phrases, sentiment, etc.)")
    custom_terms: Optional[str] = Field(
        default=None, description="Comma-separated list of custom terms for improved recognition"
    )
    language_candidates: Optional[str] = Field(
        default=None, description="Comma-separated list of language codes for multi-language support (e.g., 'en-US,nl-NL')"
    )


class JobResponse(BaseModel):
    """Response model for job submission."""

    job_id: str
    status: JobStatus
    message: str


class JobStatusResponse(BaseModel):
    """Response model for job status."""

    job_id: str
    status: JobStatus
    progress: Optional[str] = None
    result: Optional[Dict[str, Any]] = None
    error: Optional[str] = None
    created_at: str
    updated_at: str
    started_at: Optional[str] = None
    filename: Optional[str] = None
    method: Optional[str] = None
    pipeline_stages: Optional[Dict[str, Any]] = None


@app.get("/")
async def root():
    """Root endpoint."""
    return {"service": "Meeting Audio Transcription API", "version": "1.0.0", "status": "active"}


@app.get("/health")
async def health_check():
    """Health check endpoint."""
    return {"status": "healthy", "timestamp": datetime.now(timezone.utc).isoformat()}


@app.post("/api/transcribe", response_model=JobResponse)
async def transcribe_audio(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(..., description="Audio file to transcribe"),
    method: str = Form(default="azure"),
    language: Optional[str] = Form(default=None),
    enable_diarization: bool = Form(default=True),
    chunk_size: Optional[int] = Form(default=None),
    whisper_model: str = Form(default="base"),
    enable_nlp: bool = Form(default=True),
    custom_terms: Optional[str] = Form(default=None),
    language_candidates: Optional[str] = Form(default=None),
    terms_file: Optional[UploadFile] = File(default=None, description="Optional file with custom terms (one per line)"),
    # --- Advanced: Azure Speech settings ---
    profanity_filter: Optional[str] = Form(default=None),
    max_speakers: Optional[int] = Form(default=None),
    word_level_timestamps: bool = Form(default=False),
    # --- Advanced: Whisper settings ---
    whisper_temperature: Optional[float] = Form(default=None),
    whisper_prompt: Optional[str] = Form(default=None),
    # --- Advanced: HuggingFace Wav2Vec 2.0 settings ---
    hf_model: str = Form(default="facebook/wav2vec2-base-960h"),
    hf_use_api: bool = Form(default=True),
    hf_endpoint: Optional[str] = Form(default=None),
    # --- Advanced: NLP settings ---
    summary_sentence_count: Optional[int] = Form(default=None),
    nlp_features: Optional[str] = Form(default=None),
    sentiment_confidence_threshold: Optional[float] = Form(default=0.6, ge=0.0, le=1.0),
    # --- Audio pre-processing settings ---
    audio_channels: int = Form(default=1, ge=1, le=2),
    audio_sample_rate: int = Form(default=16000),
    audio_bit_rate: str = Form(default="16k"),
):
    """
    Upload an audio file and start transcription.

    Returns a job ID for tracking progress.
    """
    # Generate job ID
    job_id = str(uuid.uuid4())

    # Parse custom terms from text input or file
    terms_list = []
    if custom_terms:
        # Split by comma or newline and clean up
        terms_list = [term.strip() for term in custom_terms.replace("\n", ",").split(",") if term.strip()]

    # If terms file is uploaded, read and parse it
    if terms_file:
        try:
            terms_content = await terms_file.read()
            terms_text = terms_content.decode("utf-8")
            file_terms = [term.strip() for term in terms_text.split("\n") if term.strip()]
            terms_list.extend(file_terms)
        except Exception as e:
            logger.warning(f"Failed to read terms file: {e}")

    # Parse language candidates
    lang_candidates_list = []
    if language_candidates:
        lang_candidates_list = [lang.strip() for lang in language_candidates.split(",") if lang.strip()]

    # Save uploaded audio file
    file_path = AUDIO_DIR / f"{job_id}_{file.filename}"
    try:
        with open(file_path, "wb") as f:
            content = await file.read()
            f.write(content)
    except Exception as e:
        logger.error(f"Failed to save uploaded file: {e}")
        raise HTTPException(status_code=500, detail="Failed to save uploaded file")

    # Create job record
    jobs_db[job_id] = {
        "job_id": job_id,
        "status": JobStatus.PENDING,
        "file_path": str(file_path),
        "filename": file.filename,
        "method": method,
        "language": language,
        "enable_diarization": enable_diarization,
        "chunk_size": chunk_size,
        "whisper_model": whisper_model,
        "enable_nlp": enable_nlp,
        "custom_terms": terms_list,
        "language_candidates": lang_candidates_list,
        "profanity_filter": profanity_filter,
        "max_speakers": max_speakers,
        "word_level_timestamps": word_level_timestamps,
        "whisper_temperature": whisper_temperature,
        "whisper_prompt": whisper_prompt,
        "hf_model": hf_model,
        "hf_use_api": hf_use_api,
        "hf_endpoint": hf_endpoint,
        "summary_sentence_count": summary_sentence_count,
        "nlp_features": nlp_features,
        "sentiment_confidence_threshold": sentiment_confidence_threshold,
        "audio_channels": audio_channels,
        "audio_sample_rate": audio_sample_rate,
        "audio_bit_rate": audio_bit_rate,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "updated_at": datetime.now(timezone.utc).isoformat(),
        "result": None,
        "error": None,
    }

    # Schedule background task
    background_tasks.add_task(
        process_transcription,
        job_id=job_id,
        file_path=str(file_path),
        method=method,
        language=language,
        enable_diarization=enable_diarization,
        chunk_size=chunk_size,
        whisper_model=whisper_model,
        enable_nlp=enable_nlp,
        custom_terms=terms_list,
        language_candidates=lang_candidates_list,
        profanity_filter=profanity_filter,
        max_speakers=max_speakers,
        word_level_timestamps=word_level_timestamps,
        whisper_temperature=whisper_temperature,
        whisper_prompt=whisper_prompt,
        hf_model=hf_model,
        hf_use_api=hf_use_api,
        hf_endpoint=hf_endpoint,
        summary_sentence_count=summary_sentence_count,
        nlp_features=nlp_features,
        sentiment_confidence_threshold=sentiment_confidence_threshold,
        audio_channels=audio_channels,
        audio_sample_rate=audio_sample_rate,
        audio_bit_rate=audio_bit_rate,
    )

    return JobResponse(job_id=job_id, status=JobStatus.PENDING, message="Transcription job started")


@app.get("/api/jobs/{job_id}", response_model=JobStatusResponse)
async def get_job_status(job_id: str):
    """
    Get the status of a transcription job.
    """
    if job_id not in jobs_db:
        raise HTTPException(status_code=404, detail="Job not found")

    job = jobs_db[job_id]
    return JobStatusResponse(
        job_id=job["job_id"],
        status=job["status"],
        progress=job.get("progress"),
        result=job.get("result"),
        error=job.get("error"),
        created_at=job["created_at"],
        updated_at=job["updated_at"],
        started_at=job.get("started_at"),
        filename=job.get("filename"),
        method=job.get("method"),
        pipeline_stages=job.get("pipeline_stages"),
    )


@app.get("/api/jobs")
async def list_jobs():
    """
    List all transcription jobs.
    """
    return {
        "jobs": [
            {
                "job_id": job["job_id"],
                "status": job["status"],
                "filename": job["filename"],
                "method": job["method"],
                "created_at": job["created_at"],
            }
            for job in jobs_db.values()
        ]
    }


@app.delete("/api/jobs/{job_id}")
async def delete_job(job_id: str):
    """
    Delete a transcription job and its associated files.
    """
    if job_id not in jobs_db:
        raise HTTPException(status_code=404, detail="Job not found")

    job = jobs_db[job_id]

    # Clean up files
    try:
        file_path = Path(job["file_path"])
        if file_path.exists():
            file_path.unlink()
    except Exception as e:
        logger.warning(f"Failed to delete file: {e}")

    # Remove from database
    del jobs_db[job_id]

    return {"message": "Job deleted successfully"}


@app.post("/api/batch")
async def batch_transcribe(
    background_tasks: BackgroundTasks,
    files: List[UploadFile] = File(..., description="Audio files to transcribe (one or more)"),
    method: str = Form(default="azure"),
    language: Optional[str] = Form(default=None),
    enable_diarization: bool = Form(default=True),
    chunk_size: Optional[int] = Form(default=None),
    whisper_model: str = Form(default="base"),
    enable_nlp: bool = Form(default=True),
    custom_terms: Optional[str] = Form(default=None),
    language_candidates: Optional[str] = Form(default=None),
    parallel_batch: bool = Form(default=False),
    max_concurrent: int = Form(default=2, ge=1, le=10),
    # --- Advanced: Azure Speech settings ---
    profanity_filter: Optional[str] = Form(default=None),
    max_speakers: Optional[int] = Form(default=None),
    word_level_timestamps: bool = Form(default=False),
    # --- Advanced: Whisper settings ---
    whisper_temperature: Optional[float] = Form(default=None),
    whisper_prompt: Optional[str] = Form(default=None),
    # --- Advanced: HuggingFace Wav2Vec 2.0 settings ---
    hf_model: str = Form(default="facebook/wav2vec2-base-960h"),
    hf_use_api: bool = Form(default=True),
    hf_endpoint: Optional[str] = Form(default=None),
    # --- Advanced: NLP settings ---
    summary_sentence_count: Optional[int] = Form(default=None),
    nlp_features: Optional[str] = Form(default=None),
    sentiment_confidence_threshold: Optional[float] = Form(default=0.6, ge=0.0, le=1.0),
    # --- Audio pre-processing settings ---
    audio_channels: int = Form(default=1, ge=1, le=2),
    audio_sample_rate: int = Form(default=16000),
    audio_bit_rate: str = Form(default="16k"),
):
    """
    Upload multiple audio files and start batch transcription.

    Returns a list of job IDs for tracking progress.
    Supports parallel processing of files when parallel_batch=True.
    """
    if not files:
        raise HTTPException(status_code=422, detail="At least one file is required")

    terms_list = []
    if custom_terms:
        terms_list = [term.strip() for term in custom_terms.replace("\n", ",").split(",") if term.strip()]

    lang_candidates_list = []
    if language_candidates:
        lang_candidates_list = [lang.strip() for lang in language_candidates.split(",") if lang.strip()]

    effective_concurrent = max_concurrent if parallel_batch else 1
    semaphore = threading.Semaphore(effective_concurrent)

    job_ids = []
    for upload_file in files:
        job_id = str(uuid.uuid4())
        file_path = AUDIO_DIR / f"{job_id}_{upload_file.filename}"
        try:
            content = await upload_file.read()
            with open(file_path, "wb") as f:
                f.write(content)
        except Exception as e:
            logger.error(f"Failed to save uploaded file {upload_file.filename}: {e}")
            raise HTTPException(status_code=500, detail=f"Failed to save uploaded file: {upload_file.filename}")

        jobs_db[job_id] = {
            "job_id": job_id,
            "status": JobStatus.PENDING,
            "file_path": str(file_path),
            "filename": upload_file.filename,
            "method": method,
            "language": language,
            "enable_diarization": enable_diarization,
            "chunk_size": chunk_size,
            "whisper_model": whisper_model,
            "enable_nlp": enable_nlp,
            "custom_terms": terms_list,
            "language_candidates": lang_candidates_list,
            "profanity_filter": profanity_filter,
            "max_speakers": max_speakers,
            "word_level_timestamps": word_level_timestamps,
            "whisper_temperature": whisper_temperature,
            "whisper_prompt": whisper_prompt,
            "hf_model": hf_model,
            "hf_use_api": hf_use_api,
            "hf_endpoint": hf_endpoint,
            "summary_sentence_count": summary_sentence_count,
            "nlp_features": nlp_features,
            "sentiment_confidence_threshold": sentiment_confidence_threshold,
            "audio_channels": audio_channels,
            "audio_sample_rate": audio_sample_rate,
            "audio_bit_rate": audio_bit_rate,
            "batch_parallel": parallel_batch,
            "batch_max_concurrent": effective_concurrent,
            "created_at": datetime.now(timezone.utc).isoformat(),
            "updated_at": datetime.now(timezone.utc).isoformat(),
            "result": None,
            "error": None,
        }

        def _throttled_task(
            _jid=job_id,
            _fp=str(file_path),
            _sem=semaphore,
        ):
            with _sem:
                process_transcription(
                    job_id=_jid,
                    file_path=_fp,
                    method=method,
                    language=language,
                    enable_diarization=enable_diarization,
                    chunk_size=chunk_size,
                    whisper_model=whisper_model,
                    enable_nlp=enable_nlp,
                    custom_terms=terms_list,
                    language_candidates=lang_candidates_list,
                    profanity_filter=profanity_filter,
                    max_speakers=max_speakers,
                    word_level_timestamps=word_level_timestamps,
                    whisper_temperature=whisper_temperature,
                    whisper_prompt=whisper_prompt,
                    hf_model=hf_model,
                    hf_use_api=hf_use_api,
                    hf_endpoint=hf_endpoint,
                    summary_sentence_count=summary_sentence_count,
                    nlp_features=nlp_features,
                    sentiment_confidence_threshold=sentiment_confidence_threshold,
                    audio_channels=audio_channels,
                    audio_sample_rate=audio_sample_rate,
                    audio_bit_rate=audio_bit_rate,
                )

        background_tasks.add_task(_throttled_task)
        job_ids.append(job_id)

    return {
        "job_ids": job_ids,
        "parallel_batch": parallel_batch,
        "max_concurrent": effective_concurrent,
        "message": f"Batch of {len(job_ids)} job(s) started",
    }


def process_transcription(
    job_id: str,
    file_path: str,
    method: str,
    language: Optional[str],
    enable_diarization: bool,
    chunk_size: Optional[int],
    whisper_model: str,
    enable_nlp: bool,
    custom_terms: Optional[List[str]] = None,
    language_candidates: Optional[List[str]] = None,
    profanity_filter: Optional[str] = None,
    max_speakers: Optional[int] = None,
    word_level_timestamps: bool = False,
    whisper_temperature: Optional[float] = None,
    whisper_prompt: Optional[str] = None,
    hf_model: str = "facebook/wav2vec2-base-960h",
    hf_use_api: bool = True,
    hf_endpoint: Optional[str] = None,
    summary_sentence_count: Optional[int] = None,
    nlp_features: Optional[str] = None,
    sentiment_confidence_threshold: Optional[float] = 0.6,
    audio_channels: int = 1,
    audio_sample_rate: int = 16000,
    audio_bit_rate: str = "16k",
):
    """
    Background task to process transcription.

    Pipeline stages are reported as structured progress so the frontend
    can render a multi-stage pipeline view.  When diarization + NLP are
    both enabled they run in parallel after transcription completes.
    NLP sub-tasks also run in parallel internally.
    """

    processed_path = None

    # ------------------------------------------------------------------
    # Helper: update structured pipeline progress
    # ------------------------------------------------------------------
    def _update_pipeline(stages: Dict[str, Dict[str, Any]], progress_text: str = None):
        """Persist pipeline stages + human-readable progress."""
        fields = {
            "pipeline_stages": stages,
            "updated_at": datetime.now(timezone.utc).isoformat(),
        }
        if progress_text is not None:
            fields["progress"] = progress_text
        jobs_db.update_fields(job_id, fields)

    def _stage(status: str, detail: str = "", progress: int = 0, sub_tasks: Dict[str, str] = None):
        stage = {"status": status, "detail": detail, "progress": progress}
        if sub_tasks is not None:
            stage["sub_tasks"] = sub_tasks
        return stage

    try:
        # ------------------------------------------------------------------
        # Initialise pipeline stages
        # ------------------------------------------------------------------
        stages: Dict[str, Dict[str, Any]] = {
            "preprocessing": _stage("pending", "Waiting"),
            "transcription": _stage("pending", "Waiting"),
        }
        wants_diarization = enable_diarization and method == "whisper_api"
        if enable_diarization:
            stages["diarization"] = _stage("pending", "Waiting")
        if enable_nlp:
            stages["nlp"] = _stage("pending", "Waiting")

        jobs_db.update_fields(
            job_id,
            {
                "status": JobStatus.PROCESSING,
                "started_at": datetime.now(timezone.utc).isoformat(),
            },
        )
        _update_pipeline(stages, "Starting pipeline...")

        # ------------------------------------------------------------------
        # 1. Preprocess audio
        # ------------------------------------------------------------------
        stages["preprocessing"] = _stage("running", "Normalizing audio...", 0)
        _update_pipeline(stages, "Preprocessing audio...")

        config = ConfigManager()
        azure_config = config.get_azure_config()
        processing_config = config.get_processing_config()

        # Validate pre-processing settings
        valid_sample_rates = {8000, 16000, 22050, 44100, 48000}
        valid_bit_rates = {"16k", "32k", "64k", "128k", "192k", "256k"}
        safe_channels = max(1, min(2, audio_channels))
        safe_sample_rate = audio_sample_rate if audio_sample_rate in valid_sample_rates else 16000
        safe_bit_rate = audio_bit_rate if audio_bit_rate in valid_bit_rates else "16k"

        preprocessor = AudioPreprocessor(
            sample_rate=safe_sample_rate,
            channels=safe_channels,
            bit_rate=safe_bit_rate,
        )
        processed_path = preprocessor.normalize_audio(file_path)

        stages["preprocessing"] = _stage("done", "Audio ready", 100)
        _update_pipeline(stages, "Audio preprocessed")

        # ------------------------------------------------------------------
        # 2. Transcription
        # ------------------------------------------------------------------
        stages["transcription"] = _stage("running", "Transcribing audio...", 0)
        _update_pipeline(stages, "Transcribing audio...")

        def on_segment_recognized(segment_count):
            stages["transcription"] = _stage(
                "running",
                f"{segment_count} segment{'s' if segment_count != 1 else ''} recognized",
                min(95, segment_count * 2),
            )
            _update_pipeline(stages, f"Transcribing... ({segment_count} segments)")

        if method == "azure":
            transcriber = AzureSpeechTranscriber(
                speech_region=azure_config.speech_region,
                language=language or processing_config.default_language,
                enable_diarization=enable_diarization,
                custom_terms=custom_terms,
                language_candidates=language_candidates,
                use_managed_identity=True,
                speech_resource_id=azure_config.speech_resource_id,
                speech_endpoint=azure_config.speech_endpoint,
                profanity_filter=profanity_filter,
                max_speakers=max_speakers,
                word_level_timestamps=word_level_timestamps,
            )
            transcription_result = transcriber.transcribe_audio(
                processed_path,
                progress_callback=on_segment_recognized,
            )

        elif method == "whisper_local":
            transcriber = WhisperTranscriber(
                model_size=whisper_model,
                language=language,
                use_api=False,
                custom_terms=custom_terms,
            )
            transcription_result = transcriber.transcribe_audio(
                processed_path,
                enable_diarization=enable_diarization,
                chunk_size=chunk_size,
            )

        elif method == "whisper_api":
            if not azure_config.openai_endpoint:
                raise ValueError("Azure OpenAI endpoint not configured. Deploy Whisper via Azure AI Foundry.")
            transcriber = WhisperTranscriber(
                language=language,
                use_api=True,
                custom_terms=custom_terms,
                azure_openai_endpoint=azure_config.openai_endpoint,
                azure_openai_deployment=azure_config.openai_whisper_deployment or "whisper",
                use_managed_identity=True,
                temperature=whisper_temperature,
                initial_prompt=whisper_prompt,
            )
            stages["transcription"] = _stage(
                "running",
                "Sending audio to Azure Whisper API...",
                10,
            )
            _update_pipeline(stages, "Transcribing with Azure Whisper...")
            transcription_result = transcriber.transcribe_audio(processed_path)
        elif method == "huggingface":
            transcriber = HuggingFaceTranscriber(
                model_name=hf_model,
                language=language,
                use_api=hf_use_api,
                endpoint_url=hf_endpoint,
                custom_terms=custom_terms,
            )
            stages["transcription"] = _stage(
                "running",
                f"Transcribing with Wav2Vec 2.0 ({hf_model})...",
                10,
            )
            _update_pipeline(stages, "Transcribing with HuggingFace Wav2Vec 2.0...")
            transcription_result = transcriber.transcribe_audio(processed_path)
        else:
            raise ValueError(f"Unknown transcription method: {method}")

        stages["transcription"] = _stage(
            "done",
            f"{len(transcription_result.segments)} segments",
            100,
        )
        _update_pipeline(stages, "Transcription complete")

        # ------------------------------------------------------------------
        # 3. Parallel phase: Diarization + NLP (independent, run together)
        # ------------------------------------------------------------------
        result: Dict[str, Any] = {"transcription": transcription_result.to_dict()}

        # Build NLP options once (used by NLP task)
        nlp_opts: Dict[str, Any] = {}
        if summary_sentence_count:
            nlp_opts["summary_sentences"] = summary_sentence_count
        if nlp_features:
            features = [f.strip().lower() for f in nlp_features.split(",")]
            nlp_opts["enable_sentiment"] = "sentiment" in features
            nlp_opts["enable_key_phrases"] = "key_phrases" in features
            nlp_opts["enable_entities"] = "entities" in features
            nlp_opts["enable_action_items"] = "action_items" in features
            nlp_opts["enable_summary"] = "summary" in features
            nlp_opts["per_segment_sentiment"] = "segment_sentiment" in features
        if sentiment_confidence_threshold is not None:
            nlp_opts["sentiment_confidence_threshold"] = sentiment_confidence_threshold

        segments_dicts = None
        if transcription_result.segments:
            segments_dicts = [
                {
                    "text": seg.text,
                    "start": seg.start_time,
                    "end": seg.end_time,
                    "speaker": getattr(seg, "speaker_id", None) or "Unknown",
                }
                for seg in transcription_result.segments
            ]

        # --- Define the two parallel tasks ---
        def _run_diarization():
            """Hybrid diarization pass (Whisper → Azure Speech merge)."""
            diar_sub = {"fast_api": "running", "merge": "pending"}
            stages["diarization"] = _stage("running", "Calling Fast Transcription API...", 0, sub_tasks=diar_sub)
            _update_pipeline(stages, "Running diarization & NLP in parallel...")

            def _diar_fast_progress(c):
                # Fast API returns all at once; show completion count
                if c == 0:
                    stages["diarization"] = _stage(
                        "running",
                        "Sending audio to API...",
                        10,
                        sub_tasks=dict(diar_sub),
                    )
                else:
                    stages["diarization"] = _stage(
                        "running",
                        f"API returned {c} phrases",
                        80,
                        sub_tasks=dict(diar_sub),
                    )
                _update_pipeline(stages)

            diarizer = AzureSpeechTranscriber(
                speech_region=azure_config.speech_region,
                language=language or processing_config.default_language,
                enable_diarization=True,
                use_managed_identity=True,
                speech_resource_id=azure_config.speech_resource_id,
                speech_endpoint=azure_config.speech_endpoint,
                max_speakers=max_speakers,
            )

            # Try fast REST API first, fall back to real-time SDK
            try:
                diar_segs = diarizer.diarize_fast(
                    processed_path,
                    progress_callback=_diar_fast_progress,
                )
            except Exception as fast_err:
                logger.warning(f"Fast diarization failed, falling back to real-time: {fast_err}")
                diar_sub["fast_api"] = "error"
                diar_sub["realtime_fallback"] = "running"
                stages["diarization"] = _stage(
                    "running",
                    "Falling back to real-time...",
                    5,
                    sub_tasks=dict(diar_sub),
                )
                _update_pipeline(stages)

                def _slow_progress(c):
                    total_seg = len(transcription_result.segments)
                    pct = min(95, int(c / max(total_seg, 1) * 100))
                    diar_sub["realtime_fallback"] = "running"
                    stages["diarization"] = _stage(
                        "running",
                        f"{c} of {total_seg} segments",
                        pct,
                        sub_tasks=dict(diar_sub),
                    )
                    _update_pipeline(stages)

                diar_segs = diarizer.diarize_only(
                    processed_path,
                    progress_callback=_slow_progress,
                )
                diar_sub["realtime_fallback"] = "done"

            diar_sub["fast_api"] = diar_sub.get("fast_api", "running")
            if diar_sub["fast_api"] == "running":
                diar_sub["fast_api"] = "done"
            diar_sub["merge"] = "done"
            speakers = set(s.get("speaker_id") for s in diar_segs if s.get("speaker_id"))
            stages["diarization"] = _stage(
                "done",
                f"{len(speakers)} speakers, {len(diar_segs)} phrases",
                100,
                sub_tasks=dict(diar_sub),
            )
            _update_pipeline(stages)
            return diar_segs

        def _run_nlp():
            """NLP analysis (sub-tasks also run in parallel internally)."""
            nlp_sub = {}  # track each sub-task status
            stages["nlp"] = _stage("running", "Starting...", 0, sub_tasks=nlp_sub)
            _update_pipeline(stages, "Analyzing content...")

            nlp_subtasks_total = (
                sum(
                    [
                        (
                            nlp_opts.get("enable_key_phrases", True)
                            if not nlp_features
                            else "key_phrases" in (nlp_features or "")
                        ),
                        nlp_opts.get("enable_sentiment", True) if not nlp_features else "sentiment" in (nlp_features or ""),
                        nlp_opts.get("enable_entities", True) if not nlp_features else "entities" in (nlp_features or ""),
                        nlp_opts.get("enable_summary", True) if not nlp_features else "summary" in (nlp_features or ""),
                        (
                            nlp_opts.get("enable_action_items", True)
                            if not nlp_features
                            else "action_items" in (nlp_features or "")
                        ),
                        1,  # segment sentiment
                    ]
                )
                or 6
            )
            completed_count = [0]

            def _nlp_progress(task_name, status):
                nlp_sub[task_name] = status
                if status == "done":
                    completed_count[0] += 1
                    pct = int(completed_count[0] / nlp_subtasks_total * 100)
                    stages["nlp"] = _stage(
                        "running",
                        f"{completed_count[0]}/{nlp_subtasks_total} tasks done",
                        min(95, pct),
                        sub_tasks=dict(nlp_sub),
                    )
                elif status == "running":
                    stages["nlp"]["detail"] = f"Running {task_name}..."
                    stages["nlp"]["sub_tasks"] = dict(nlp_sub)
                _update_pipeline(stages)

            analyzer = ContentAnalyzer(
                text_analytics_endpoint=azure_config.text_analytics_endpoint,
                use_managed_identity=True,
            )
            nlp_result = analyzer.analyze_transcription(
                transcription_result.full_text,
                segments=segments_dicts,
                nlp_options=nlp_opts if nlp_opts else None,
                progress_callback=_nlp_progress,
            )
            stages["nlp"] = _stage("done", "Analysis complete", 100)
            _update_pipeline(stages)
            return nlp_result

        # --- Launch parallel tasks ---
        parallel_futures: Dict[str, Any] = {}
        with ThreadPoolExecutor(max_workers=2, thread_name_prefix="pipeline") as pool:
            if wants_diarization:
                parallel_futures["diarization"] = pool.submit(_run_diarization)
            elif enable_diarization:
                # Azure method already did diarization inline
                stages.get("diarization", {}).update(
                    {"status": "done", "detail": "Inline with transcription", "progress": 100}
                )

            if enable_nlp and transcription_result.full_text:
                parallel_futures["nlp"] = pool.submit(_run_nlp)

            if parallel_futures:
                active_names = " & ".join(k.title() for k in parallel_futures)
                _update_pipeline(stages, f"Running {active_names} in parallel...")

            # Wait for all to finish
            for key, future in parallel_futures.items():
                try:
                    task_result = future.result()
                    if key == "diarization":
                        transcription_result = WhisperTranscriber.merge_diarization(
                            transcription_result,
                            task_result,
                        )
                        result["transcription"] = transcription_result.to_dict()
                        logger.info(f"Hybrid diarization merged for job {job_id}")
                    elif key == "nlp":
                        result["nlp_analysis"] = task_result.to_dict()
                except Exception as e:
                    logger.warning(f"Parallel task '{key}' failed: {e}")
                    if key in stages:
                        stages[key] = _stage("error", str(e)[:120], 0)

        # If NLP ran but diarization also ran, re-build segment sentiments
        # with updated speaker IDs (from diarization merge)
        if wants_diarization and enable_nlp and "nlp_analysis" in result and "diarization" in parallel_futures:
            # Update segment speaker info in NLP results
            if segments_dicts and result.get("nlp_analysis", {}).get("segment_sentiments"):
                updated_segs = result["transcription"].get("segments", [])
                for ss in result["nlp_analysis"]["segment_sentiments"]:
                    idx = ss.get("index", -1)
                    if 0 <= idx < len(updated_segs):
                        ss["speaker"] = updated_segs[idx].get("speaker_id") or ss["speaker"]

        # ------------------------------------------------------------------
        # 4. Complete
        # ------------------------------------------------------------------
        for k in stages:
            if stages[k]["status"] not in ("done", "error"):
                stages[k] = _stage("done", "Skipped", 100)
        _update_pipeline(stages, "Completed")

        jobs_db.update_fields(
            job_id,
            {
                "status": JobStatus.COMPLETED,
                "result": result,
                "error": None,
                "progress": "Completed",
                "updated_at": datetime.now(timezone.utc).isoformat(),
            },
        )
        logger.info(f"Job {job_id} completed successfully")

    except Exception as e:
        logger.error(f"Job {job_id} failed: {e}", exc_info=True)
        jobs_db.update_fields(
            job_id,
            {
                "status": JobStatus.FAILED,
                "error": str(e),
                "updated_at": datetime.now(timezone.utc).isoformat(),
            },
        )

    finally:
        try:
            if processed_path and os.path.exists(processed_path):
                os.unlink(processed_path)
        except Exception as e:
            logger.warning(f"Failed to clean up processed file: {e}")


@app.get("/api/audio/{job_id}")
async def serve_audio(job_id: str, request: Request):
    """
    Serve the audio file for a completed job with HTTP Range support
    so the browser can seek within the audio.
    """
    if job_id not in jobs_db:
        raise HTTPException(status_code=404, detail="Job not found")

    job = jobs_db[job_id]
    file_path = Path(job["file_path"])

    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Audio file not found")

    import mimetypes

    mime_type, _ = mimetypes.guess_type(job["filename"])
    if not mime_type or not mime_type.startswith("audio/"):
        mime_type = "application/octet-stream"

    file_size = file_path.stat().st_size
    range_header = request.headers.get("range")

    if range_header:
        # Parse Range header (e.g. "bytes=12345-" or "bytes=12345-67890")
        range_spec = range_header.replace("bytes=", "").strip()
        parts = range_spec.split("-")
        start = int(parts[0]) if parts[0] else 0
        end = int(parts[1]) if parts[1] else file_size - 1
        end = min(end, file_size - 1)
        length = end - start + 1

        def _range_iter():
            with open(file_path, "rb") as f:
                f.seek(start)
                remaining = length
                while remaining > 0:
                    chunk = f.read(min(8192, remaining))
                    if not chunk:
                        break
                    remaining -= len(chunk)
                    yield chunk

        return StreamingResponse(
            _range_iter(),
            status_code=206,
            headers={
                "Content-Range": f"bytes {start}-{end}/{file_size}",
                "Accept-Ranges": "bytes",
                "Content-Length": str(length),
                "Content-Type": mime_type,
                "Content-Disposition": f'inline; filename="{job["filename"]}"',
            },
            media_type=mime_type,
        )
    else:
        # Full file response with Accept-Ranges header
        from fastapi.responses import FileResponse

        return FileResponse(
            path=str(file_path),
            media_type=mime_type,
            filename=job["filename"],
            headers={"Accept-Ranges": "bytes"},
        )


@app.post("/api/export/{job_id}")
async def export_transcription(job_id: str, format: str = Form(...)):
    """
    Export transcription in different formats (txt, docx, pdf).
    """
    if job_id not in jobs_db:
        raise HTTPException(status_code=404, detail="Job not found")

    job = jobs_db[job_id]

    if job["status"] != JobStatus.COMPLETED:
        raise HTTPException(status_code=400, detail="Job not completed yet")

    if not job.get("result"):
        raise HTTPException(status_code=404, detail="No transcription result found")

    transcription = job["result"]["transcription"]
    nlp_analysis = job["result"].get("nlp_analysis")

    try:
        if format == "txt":
            return export_as_txt(transcription, nlp_analysis, job["filename"])
        elif format == "docx":
            return export_as_docx(transcription, nlp_analysis, job["filename"])
        elif format == "pdf":
            return export_as_pdf(transcription, nlp_analysis, job["filename"])
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported format: {format}")
    except Exception as e:
        logger.error(f"Export failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Export failed: {str(e)}")


def export_as_txt(transcription: Dict[str, Any], nlp_analysis: Optional[Dict[str, Any]], filename: str):
    """Export transcription as plain text file."""
    from fastapi.responses import Response
    from io import StringIO

    output = StringIO()
    output.write(f"Transcription: {filename}\n")
    output.write("=" * 80 + "\n\n")

    # Metadata
    if transcription.get("language"):
        output.write(f"Language: {transcription['language']}\n")
    if transcription.get("duration"):
        output.write(f"Duration: {transcription['duration']:.2f} seconds\n")
    if transcription.get("metadata", {}).get("speaker_count"):
        output.write(f"Speakers: {transcription['metadata']['speaker_count']}\n")
    output.write("\n")

    # Full text
    output.write("Full Transcription:\n")
    output.write("-" * 80 + "\n")
    output.write(transcription.get("full_text", "") + "\n\n")

    # Segments with timestamps
    if transcription.get("segments"):
        output.write("\nDetailed Segments:\n")
        output.write("-" * 80 + "\n")
        for segment in transcription["segments"]:
            timestamp = f"[{segment['start_time']:.1f}s - {segment['end_time']:.1f}s]"
            speaker = f"{segment.get('speaker_id', 'Unknown')}: " if segment.get("speaker_id") else ""
            output.write(f"{timestamp} {speaker}{segment['text']}\n")

    # NLP Analysis
    if nlp_analysis:
        output.write("\n\nContent Analysis:\n")
        output.write("=" * 80 + "\n")

        if nlp_analysis.get("sentiment"):
            output.write(f"\nSentiment: {nlp_analysis['sentiment'].get('overall', 'N/A')}\n")

        if nlp_analysis.get("key_phrases"):
            output.write("\nKey Phrases:\n")
            for phrase in nlp_analysis["key_phrases"][:MAX_KEY_PHRASES_EXPORT]:
                output.write(f"  - {phrase['text']}\n")

    content = output.getvalue()
    return Response(
        content=content,
        media_type="text/plain",
        headers={"Content-Disposition": f"attachment; filename={filename.rsplit('.', 1)[0]}.txt"},
    )


def export_as_docx(transcription: Dict[str, Any], nlp_analysis: Optional[Dict[str, Any]], filename: str):
    """Export transcription as Word document."""
    from docx import Document
    from docx.shared import RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from io import BytesIO
    from fastapi.responses import Response

    doc = Document()

    # Title
    title = doc.add_heading(f"Transcription: {filename}", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Metadata section
    doc.add_heading("Metadata", level=1)
    metadata_para = doc.add_paragraph()
    if transcription.get("language"):
        metadata_para.add_run("Language: ").bold = True
        metadata_para.add_run(f"{transcription['language']}\n")
    if transcription.get("duration"):
        metadata_para.add_run("Duration: ").bold = True
        metadata_para.add_run(f"{transcription['duration']:.2f} seconds\n")
    if transcription.get("metadata", {}).get("speaker_count"):
        metadata_para.add_run("Speakers: ").bold = True
        metadata_para.add_run(f"{transcription['metadata']['speaker_count']}\n")

    # Full transcription
    doc.add_heading("Full Transcription", level=1)
    doc.add_paragraph(transcription.get("full_text", ""))

    # Segments
    if transcription.get("segments"):
        doc.add_page_break()
        doc.add_heading("Detailed Segments", level=1)
        for segment in transcription["segments"]:
            para = doc.add_paragraph()
            # Timestamp
            timestamp_run = para.add_run(f"[{segment['start_time']:.1f}s - {segment['end_time']:.1f}s] ")
            timestamp_run.font.color.rgb = RGBColor(0, 102, 204)
            timestamp_run.bold = True
            # Speaker
            if segment.get("speaker_id"):
                speaker_run = para.add_run(f"{segment['speaker_id']}: ")
                speaker_run.font.color.rgb = RGBColor(155, 89, 182)
                speaker_run.bold = True
            # Text
            para.add_run(segment["text"])

    # NLP Analysis
    if nlp_analysis:
        doc.add_page_break()
        doc.add_heading("Content Analysis", level=1)

        if nlp_analysis.get("sentiment"):
            doc.add_heading("Sentiment", level=2)
            doc.add_paragraph(f"Overall: {nlp_analysis['sentiment'].get('overall', 'N/A')}")

        if nlp_analysis.get("key_phrases"):
            doc.add_heading("Key Phrases", level=2)
            for phrase in nlp_analysis["key_phrases"][:MAX_KEY_PHRASES_EXPORT]:
                doc.add_paragraph(phrase["text"], style="List Bullet")

    # Save to BytesIO
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)

    return Response(
        content=bio.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename.rsplit('.', 1)[0]}.docx"},
    )


def export_as_pdf(transcription: Dict[str, Any], nlp_analysis: Optional[Dict[str, Any]], filename: str):
    """Export transcription as PDF document."""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.enums import TA_CENTER
    from io import BytesIO
    from fastapi.responses import Response

    bio = BytesIO()
    doc = SimpleDocTemplate(bio, pagesize=letter, topMargin=0.5 * inch, bottomMargin=0.5 * inch)
    story = []
    styles = getSampleStyleSheet()

    # Custom styles
    title_style = ParagraphStyle(
        "CustomTitle", parent=styles["Heading1"], fontSize=24, textColor="#667eea", alignment=TA_CENTER, spaceAfter=12
    )

    heading_style = ParagraphStyle("CustomHeading", parent=styles["Heading2"], fontSize=14, textColor="#667eea", spaceAfter=8)

    # Title
    story.append(Paragraph(f"Transcription: {filename}", title_style))
    story.append(Spacer(1, 0.2 * inch))

    # Metadata
    story.append(Paragraph("Metadata", heading_style))
    metadata_text = ""
    if transcription.get("language"):
        metadata_text += f"<b>Language:</b> {transcription['language']}<br/>"
    if transcription.get("duration"):
        metadata_text += f"<b>Duration:</b> {transcription['duration']:.2f} seconds<br/>"
    if transcription.get("metadata", {}).get("speaker_count"):
        metadata_text += f"<b>Speakers:</b> {transcription['metadata']['speaker_count']}<br/>"
    story.append(Paragraph(metadata_text, styles["Normal"]))
    story.append(Spacer(1, 0.2 * inch))

    # Full transcription
    story.append(Paragraph("Full Transcription", heading_style))
    story.append(Paragraph(transcription.get("full_text", ""), styles["Normal"]))
    story.append(PageBreak())

    # Segments
    if transcription.get("segments"):
        story.append(Paragraph("Detailed Segments", heading_style))
        for segment in transcription["segments"]:
            timestamp = f"[{segment['start_time']:.1f}s - {segment['end_time']:.1f}s]"
            speaker = f"<b>{segment.get('speaker_id', 'Unknown')}:</b> " if segment.get("speaker_id") else ""
            text = f'<font color="#3498db">{timestamp}</font> {speaker}{segment["text"]}'
            story.append(Paragraph(text, styles["Normal"]))
            story.append(Spacer(1, 0.05 * inch))

    # NLP Analysis
    if nlp_analysis:
        story.append(PageBreak())
        story.append(Paragraph("Content Analysis", heading_style))

        if nlp_analysis.get("sentiment"):
            story.append(Paragraph("Sentiment", styles["Heading3"]))
            story.append(Paragraph(f"Overall: {nlp_analysis['sentiment'].get('overall', 'N/A')}", styles["Normal"]))
            story.append(Spacer(1, 0.1 * inch))

        if nlp_analysis.get("key_phrases"):
            story.append(Paragraph("Key Phrases", styles["Heading3"]))
            phrases = ", ".join([phrase["text"] for phrase in nlp_analysis["key_phrases"][:MAX_KEY_PHRASES_EXPORT]])
            story.append(Paragraph(phrases, styles["Normal"]))

    # Build PDF
    doc.build(story)
    bio.seek(0)

    return Response(
        content=bio.getvalue(),
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={filename.rsplit('.', 1)[0]}.pdf"},
    )


if __name__ == "__main__":
    import uvicorn

    # Use environment variable for host binding, default to localhost for security
    # Set API_HOST=0.0.0.0 in production environments where needed
    host = os.getenv("API_HOST", "127.0.0.1")
    port = int(os.getenv("API_PORT", "8000"))
    uvicorn.run(app, host=host, port=port)
